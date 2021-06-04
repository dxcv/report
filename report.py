import os
import pandas as pd
import pymysql
from WindPy import *
from docx import *
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


class Function:
    class Base:
        def __init__(self):
            self.data = pd.DataFrame(columns=['业务日期', '投组单元名称', '产品分类', '名称', '市值', '到期日', '起息日', '成本', '估值'])
            self.flow = pd.DataFrame(columns=['名称', '类别', '交易日', '方向', '金额', '交易投组', '对手方', '净价'])
            self.bs()
            self.prepare_stream()

        def bs(self):
            pass

        def asset(self):
            tmp = self.data[self.data['市值'] > 0].copy(deep=True)
            tmp['市值'] = abs(tmp['市值'])
            return tmp

        def loan(self):
            tmp = self.data[self.data['市值'] < 0].copy(deep=True)
            tmp['市值'] = abs(tmp['市值'])
            return tmp

        def prepare_stream(self):
            pass

        def stream(self):
            return self.flow

    class TY(Base):
        def bs(self):
            if os.path.exists("data/福建农信金融市场业务余额表（汇总）.xls"):
                tmp = pd.read_excel("data/福建农信金融市场业务余额表（汇总）.xls")
                res = pd.to_datetime(tmp.loc[0, 'Unnamed: 10'], format='%Y-%m-%d')
                tmp = pd.read_excel("data/福建农信金融市场业务余额表（汇总）.xls", header=2)
                tmp = tmp[tmp['交易余额'] > 0]
                tmp.loc[tmp['投资类别（资产/负债）'] == '负债', '交易余额'] = -tmp.loc[tmp['投资类别（资产/负债）'] == '负债', '交易余额']
                tmp = tmp[['业务类型', '产品名称', '交易余额', '到期日', '投资开始日/起息日']]
                tmp.columns = ['产品分类', '名称', '市值', '到期日', '起息日']
                tmp['投组单元名称'] = '线下'
                tmp['市值'] = tmp['市值'] / 100000000
                tmp['业务日期'] = res
                self.data = self.data.append(tmp)
                self.data = self.data.reset_index(drop=True)

            if os.path.exists("data/指定成本与FIFO损益分析-新.xls"):
                tmp = pd.read_excel("data/指定成本与FIFO损益分析-新.xls")
                tmp = tmp[~tmp['Unnamed: 4'].isna()]
                tmp = tmp[(tmp['交易投组'] != '2同业-同业投资(jy2@xmrcb)') & (tmp['交易投组'] != '自营-债基-底层资产(林文妹)')]
                tmp = tmp[~tmp['交易投组'].str.contains('资金往来')]
                tmp.loc[~tmp['债券名称'].isna(), '产品分类'] = '债券'
                tmp.loc[tmp['交易投组'].str.contains('回购'), '产品分类'] = '回购'
                tmp.loc[tmp['交易投组'].str.contains('拆借'), '产品分类'] = '拆借'
                tmp.loc[tmp['交易投组'].str.contains('同业借款'), '产品分类'] = '同业借款'
                tmp.loc[tmp['交易投组'].str.contains('债券借贷'), '产品分类'] = '债券借贷'
                tmp.loc[tmp['产品分类'].isna(), '产品分类'] = '其他'
                tmp = tmp[['交易投组', '产品分类', 'Unnamed: 4', '市值', '到期日', '起息日', '原始购入成本价', '市场净价']]
                tmp.columns = ['投组单元名称', '产品分类', '名称', '市值', '到期日', '起息日', '成本', '估值']
                tmp['市值'] = tmp['市值'] / 100000000
                tmp.loc[tmp['产品分类'] == '债券', '名称'] = tmp.loc[tmp['产品分类'] == '债券', '名称'].map(
                    lambda lambda_x: lambda_x[:lambda_x.rfind("(")])
                tmp['业务日期'] = self.data.loc[0, '业务日期']
                if not w.isconnected():
                    w.start()
                exchange = w.wss("USDCNY.EX", "close", "tradeDate=" +
                                 str(self.data.loc[0, '业务日期'].date()).replace("-", "") + ";priceAdj=U;cycle=D")
                for x in tmp[tmp['投组单元名称'].str.contains('美元')].index.tolist():
                    tmp.loc[x, '市值'] = tmp.loc[x, '市值'] * exchange.Data[0][0]
                self.data = self.data.append(tmp)
                self.data = self.data.reset_index(drop=True)

        def prepare_stream(self):
            if os.path.exists('data/交易查询与维护_现券.xls'):
                tmp = pd.read_excel('data/交易查询与维护_现券.xls', header=1)
                tmp['名称'] = tmp['债券名称'].map(lambda x: x[:x.rfind('(')])
                tmp['类别'] = '债券'
                tmp['交易日'] = tmp['交易日'].map(lambda x: pd.to_datetime(x, format='%Y-%m-%d'))
                tmp['方向'] = tmp['交易方向']
                tmp['金额'] = tmp['结算金额(元)'] / 100000000
                tmp = tmp[['名称', '类别', '交易日', '方向', '金额', '交易投组', '对手方', '净价']]
                self.flow = self.flow.append(tmp)
                self.flow = self.flow.reset_index(drop=True)

            if os.path.exists('data/交易查询与维护_质押式回购.xls'):
                tmp = pd.read_excel('data/交易查询与维护_质押式回购.xls', header=1)
                tmp['名称'] = tmp['回购名称']
                tmp['交易日'] = tmp['交易日'].map(lambda x: pd.to_datetime(x, format='%Y-%m-%d'))
                tmp['方向'] = tmp['回购方向']
                tmp['金额'] = tmp['交易金额(元)'] / 100000000
                tmp['类别'] = '回购'
                tmp = tmp[['名称', '类别', '交易日', '方向', '金额', '交易投组', '对手方']]
                self.flow = self.flow.append(tmp)
                self.flow = self.flow.reset_index(drop=True)

            if os.path.exists('data/交易查询与维护_同业拆借.xls'):
                tmp = pd.read_excel('data/交易查询与维护_同业拆借.xls', header=1)
                tmp['名称'] = tmp['交易品种']
                tmp['交易日'] = tmp['交易日'].map(lambda x: pd.to_datetime(x, format='%Y-%m-%d'))
                tmp['方向'] = tmp['拆借方向']
                tmp['金额'] = tmp['拆借金额(万)'] / 10000
                tmp['类别'] = '同业拆借'
                tmp = tmp[['名称', '类别', '交易日', '方向', '金额', '交易投组', '对手方']]
                self.flow = self.flow.append(tmp)
                self.flow = self.flow.reset_index(drop=True)

            if os.path.exists('data/交易查询与维护_债券借贷.xls'):
                tmp = pd.read_excel('data/交易查询与维护_债券借贷.xls', header=1)
                tmp['名称'] = '债券借贷'
                tmp['交易日'] = tmp['交易日'].map(lambda x: pd.to_datetime(x, format='%Y-%m-%d'))
                tmp['方向'] = tmp['交易方向']
                tmp['金额'] = tmp['标的券券面总额(万)'] / 10000
                tmp['类别'] = '债券借贷'
                tmp = tmp[['名称', '类别', '交易日', '方向', '金额', '交易投组', '对手方']]
                self.flow = self.flow.append(tmp)
                self.flow = self.flow.reset_index(drop=True)

            if os.path.exists('data/交易查询与维护_同业借款.xls'):
                tmp = pd.read_excel('data/交易查询与维护_同业借款.xls', header=1)
                tmp['名称'] = '同业借款'
                tmp['交易日'] = tmp['交易日期'].map(lambda x: pd.to_datetime(x, format='%Y-%m-%d'))
                tmp['方向'] = tmp['交易方向']
                tmp['金额'] = tmp['交易本金(万)'] / 10000
                tmp['类别'] = '同业借款'
                tmp = tmp[['名称', '类别', '交易日', '方向', '金额', '交易投组', '对手方']]
                self.flow = self.flow.append(tmp)
                self.flow = self.flow.reset_index(drop=True)

            if os.path.exists('data/交易查询与维护_承销.xls'):
                tmp = pd.read_excel('data/交易查询与维护_承销.xls')
                tmp['名称'] = tmp['债券'].map(lambda x: x[:x.rfind('(')])
                tmp['交易日'] = tmp['交易日'].map(lambda x: pd.to_datetime(x, format='%Y-%m-%d'))
                tmp['交易投组'] = tmp['投组']
                tmp['方向'] = tmp['交易类型'].map(
                    lambda x: {'承销买入': '买入', '一级市场投资': '买入', '分销入': '买入', '一级市场分销卖出': '卖出',
                               '二级市场分销卖出': '卖出', '转自营': '卖出'}.get(x))
                tmp['金额'] = tmp['缴款金额(元)'] / 100000000
                tmp['交易投组'] = tmp['投组']
                tmp['净价'] = tmp['净价(元)']
                tmp['类别'] = '债券'
                tmp = tmp[['名称', '类别', '交易日', '方向', '金额', '交易投组', '对手方', '净价']]
                self.flow = self.flow.append(tmp)
                self.flow = self.flow.reset_index(drop=True)

            if os.path.exists('data/存放同业交易明细.xls'):
                tmp = pd.read_excel('data/存放同业交易明细.xls', header=2)
                tmp['名称'] = '存放同业'
                tmp['类别'] = '存放同业'
                tmp['交易日'] = tmp['起息日期'].map(
                    lambda x: pd.to_datetime(x.replace("年", '/').replace('月', '/').replace('日', ''),
                                             format='%Y/%m/%d'))
                tmp['方向'] = tmp['交易方向']
                tmp.loc[tmp['存入金额（元）'].isna(), '金额'] = tmp.loc[tmp['存入金额（元）'].isna(), '结算金额(元)'] / 100000000
                tmp.loc[tmp['金额'].isna(), '金额'] = tmp.loc[tmp['金额'].isna(), '存入金额（元）'] / 100000000
                tmp['交易投组'] = '存放同业'
                tmp['对手方'] = tmp['交易对手']
                tmp = tmp[['名称', '类别', '交易日', '方向', '金额', '交易投组', '对手方']]
                self.flow = self.flow.append(tmp)
                self.flow = self.flow.reset_index(drop=True)

            if os.path.exists('data/同业存放交易明细表.xls'):
                tmp = pd.read_excel('data/同业存放交易明细表.xls', header=2)
                tmp['名称'] = '同业存放'
                tmp['类别'] = '同业存放'
                tmp['交易日'] = tmp['起息日期'].map(lambda x: pd.to_datetime(x, format='%Y-%m-%d'))
                tmp['方向'] = tmp['交易方向']
                tmp.loc[tmp['存入金额（元）'].isna(), '金额'] = tmp.loc[tmp['存入金额（元）'].isna(), '结算金额(元)'] / 100000000
                tmp.loc[tmp['金额'].isna(), '金额'] = tmp.loc[tmp['金额'].isna(), '存入金额（元）'] / 100000000
                tmp['交易投组'] = '同业存放'
                tmp['对手方'] = tmp['交易对手']
                tmp = tmp[['名称', '类别', '交易日', '方向', '金额', '交易投组', '对手方']]
                self.flow = self.flow.append(tmp)
                self.flow = self.flow.reset_index(drop=True)

            if os.path.exists('data/上存约期存款明细表.xls'):
                tmp = pd.read_excel('data/上存约期存款明细表.xls', header=2)
                tmp['名称'] = '上存约期存款'
                tmp['类别'] = '上存约期存款'
                tmp['交易日'] = tmp['起息日'].map(lambda x: pd.to_datetime(x, format='%Y-%m-%d'))
                tmp.loc[tmp['上存模式'].isna(), '方向'] = '支取'
                tmp['方向'].fillna({'上存模式': '上存'}, inplace=True)
                tmp.loc[tmp['方向'] == '上存', '金额'] = tmp.loc[tmp['方向'] == '上存', '上存金额(元)'] / 100000000
                tmp.loc[tmp['方向'] == '支取', '金额'] = tmp.loc[tmp['方向'] == '支取', '支取金额（元）'] / 100000000
                tmp['交易投组'] = '约期存款'
                tmp['对手方'] = '省联社'
                tmp = tmp[['名称', '类别', '交易日', '方向', '金额', '交易投组', '对手方']]
                self.flow = self.flow.append(tmp)
                self.flow = self.flow.reset_index(drop=True)

            if os.path.exists('data/委托存放交易明细表.xls'):
                tmp = pd.read_excel('data/委托存放交易明细表.xls', header=2)
                tmp['名称'] = '代理存放'
                tmp['类别'] = '代理存放'
                tmp['交易日'] = tmp['起息日期'].map(lambda x: pd.to_datetime(x, format='%Y-%m-%d'))
                tmp['方向'] = tmp['交易类型']
                tmp.loc[tmp['方向'].str.contains('到期'), '金额'] = tmp.loc[
                                                                  tmp['方向'].str.contains('到期'), '支取金额(元)'] / 100000000
                tmp.loc[tmp['金额'].isna(), '金额'] = tmp.loc[tmp['金额'].isna(), '存放金额（元）'] / 100000000
                tmp['交易投组'] = '代理存放'
                tmp['对手方'] = tmp['存放行']
                tmp = tmp[['名称', '类别', '交易日', '方向', '金额', '交易投组', '对手方']]
                self.flow = self.flow.append(tmp)
                self.flow = self.flow.reset_index(drop=True)

            if os.path.exists('data/其他投资交易明细表.xls'):
                tmp = pd.read_excel('data/其他投资交易明细表.xls', header=2)
                tmp['名称'] = tmp['资产名称']
                tmp['类别'] = '其他投资'
                tmp['交易日'] = tmp['申请日期'].map(lambda x: pd.to_datetime(x, format='%Y-%m-%d'))
                tmp['方向'] = tmp['交易方向']
                tmp['金额'] = tmp['金额（元）'] / 100000000
                tmp['交易投组'] = '其他投资'
                tmp['对手方'] = tmp['交易对手']
                tmp = tmp[['名称', '类别', '交易日', '方向', '金额', '交易投组', '对手方']]
                self.flow = self.flow.append(tmp)
                self.flow = self.flow.reset_index(drop=True)

    class LC(Base):
        def bs(self):
            if os.path.exists("data/估值余额查询.xls"):
                tmp = pd.read_excel("data/估值余额查询.xls")
                tmp = tmp[tmp['名称'] != '鑫安利得7号']
                tmp = tmp[(tmp['投组单元名称'] == '丰裕') | (tmp['投组单元名称'] == '鑫安利得7号') | (tmp['投组单元名称'] == '丰盈专属')]
                tmp.loc[tmp['产品分类2'] == '直接融资工具', '产品分类'] = '理财直融工具'
                tmp = tmp[['业务日期', '投组单元名称', '产品分类', '名称', '市值(元)', '到期日', '建仓时间']]
                tmp['投组单元名称'].replace('鑫安利得7号', '丰裕', inplace=True)
                tmp['产品分类'].replace('资产证券化债券', '债券', inplace=True)
                tmp['市值'] = tmp['市值(元)'] / 100000000
                tmp = tmp[['业务日期', '投组单元名称', '产品分类', '名称', '市值', '到期日', '建仓时间']]
                tmp = tmp[tmp['市值'] != 0].copy(deep=True)
                tmp.loc[(tmp['市值'] > 0) & (tmp['产品分类'] == '质押式回购'), '产品分类'] = '买入返售金融资产'
                tmp.loc[(tmp['市值'] < 0) & (tmp['产品分类'] == '质押式回购'), '产品分类'] = '卖出回购金融资产款'
                tmp['建仓时间'] = tmp['建仓时间'].map(lambda z: pd.to_datetime(z.split()[0], format='%Y-%m-%d'))
                tmp.columns = ['业务日期', '投组单元名称', '产品分类', '名称', '市值', '到期日', '起息日']
                tmp['name'] = tmp['投组单元名称'] + tmp['名称']
                db = pymysql.connect(host='localhost', port=3306, user='root', password='root', db='cost',
                                     charset='utf8')
                cost = pd.read_sql("select part,name,cost from licai", db)
                cost['name'] = cost['part'] + cost['name']
                tmp = pd.merge(tmp, cost, how='left', on='name')
                tmp = tmp[['业务日期', '投组单元名称', '产品分类', '名称', '市值', '到期日', '起息日', 'cost']]
                tmp.columns = ['业务日期', '投组单元名称', '产品分类', '名称', '市值', '到期日', '起息日', '成本']
                self.data = self.data.append(tmp)
                self.data = self.data.reset_index(drop=True)

            for x in os.listdir('data/'):
                if x.split('.')[0].endswith('估值报表'):
                    name = x.split("_")[0].split("-")[1]
                    fix_date = pd.to_datetime(x.split("_")[1], format='%Y%m%d')
                    tmp = pd.read_excel('data/' + x, header=4)
                    tmp = tmp[tmp.index > 1].copy(deep=True)
                    code = list(set([str(x) for x in tmp['科目代码'].tolist() if len(str(x).split('.')) == 1]))
                    code = [x for x in code if x[0] in ['0', '1', '2', '3', '4', '5', '6', '7', '8', '9']]
                    detail_code = tmp.loc[(tmp['科目名称'] == '债券投资') | (tmp['科目名称'] == '交易类资产支持证券'), '科目代码'].tolist()
                    detail = []
                    for y in detail_code:
                        detail += tmp.loc[tmp['科目代码'].map(
                            lambda z: str(z).startswith(y) and len(str(z).split('.')) == 4), '科目名称'].tolist()
                    detail = list(set(detail))
                    temp = tmp[
                        tmp['科目名称'].map(lambda z: '(总价)' in str(z) and str(z).replace('(总价)', '') in detail)].copy(
                        deep=True)
                    temp['科目名称'] = temp['科目名称'].map(lambda z: z.replace('(总价)', ''))
                    temp['产品分类'] = '债券'
                    detail = set(detail).difference(set(temp['科目名称'].tolist()))
                    temp_extend = tmp[tmp['科目名称'].map(lambda z: z in detail) & (
                        tmp['科目代码'].map(lambda z: str(z)[:4] in detail_code))].copy(deep=True)
                    temp_extend['产品分类'] = '债券'
                    tmp = tmp[tmp['科目代码'].map(lambda z: z in code)]
                    tmp = tmp[tmp['科目代码'].map(lambda z: z not in detail_code)]
                    tmp.loc[tmp['科目代码'].map(lambda z: z[0] == '2'), '市值'] = -tmp.loc[
                        tmp['科目代码'].map(lambda z: z[0] == '2'), '市值']
                    tmp['产品分类'] = tmp['科目名称']
                    temp = temp.append(tmp)
                    temp = temp.append(temp_extend)
                    temp['业务日期'] = fix_date
                    temp['投组单元名称'] = name
                    temp['名称'] = temp['科目名称']
                    temp['成本'] = temp['单位成本']
                    temp['市值'] = temp['市值'] / 100000000
                    temp = temp.append(
                        {'业务日期': fix_date,
                         '投组单元名称': name,
                         '产品分类': '理财产品',
                         '名称': '理财产品',
                         '市值': -temp['市值'].sum(),
                         '成本': float('nan')}, ignore_index=True)
                    temp = temp[['业务日期', '投组单元名称', '产品分类', '名称', '市值', '成本']]
                    self.data = self.data.append(temp)
            self.data = self.data.reset_index(drop=True)

            db = pymysql.connect(host='localhost', port=3306, user='root', password='root', db='pac', charset='utf8')
            instrument = pd.read_sql("select name from instrument_am", db)['name'].tolist()
            if os.path.exists("data/利率型.xls"):
                tmp = pd.read_excel("data/利率型.xls")
                tmp_index = tmp[~tmp['理财产品/内部投组名称'].isna()].index.tolist()
                tmp_index.append(len(tmp))
                for x in range(len(tmp_index) - 1):
                    title = tmp.loc[tmp_index[x], '理财产品/内部投组名称']
                    for line in range(tmp_index[x], tmp_index[x + 1]):
                        tmp.loc[line, '理财产品/内部投组名称'] = title
                tmp = tmp[~tmp['投资资产明细'].isna()]
                tmp = tmp[(tmp['理财产品/内部投组名称'] != '厦门农商丰裕理财计划') & (tmp['理财产品/内部投组名称'] != '厦门农商银行丰盈专属人民币理财计划')]
                tmp['市值'] = tmp['投资金额(万元)'] / 10000
                tmp['资产名称'].replace(float('nan'), "", inplace=True)
                tmp['资产名称'] = tmp['资产名称'].map(lambda lambda_x: lambda_x[:lambda_x.rfind("(")])
                tmp = tmp[['理财产品/内部投组名称', '投资资产明细', '资产名称', '市值', '到期日', '起息日', '买入价格/100元', '估值/100元']]
                tmp.columns = ['投组单元名称', '产品分类', '名称', '市值', '到期日', '起息日', '成本', '估值']
                tmp['业务日期'] = self.data.loc[0, '业务日期']
                for x in instrument:
                    tmp.loc[tmp['名称'] == x, '产品分类'] = '理财直融工具'
                self.data = self.data.append(tmp)
                self.data = self.data.reset_index(drop=True)
            if os.path.exists("data/净值型.xls"):
                tmp = pd.read_excel("data/净值型.xls")
                tmp_index = tmp[~tmp['理财产品/内部投组名称'].isna()].index.tolist()
                tmp_index.append(len(tmp))
                for x in range(len(tmp_index) - 1):
                    title = tmp.loc[tmp_index[x], '理财产品/内部投组名称']
                    for line in range(tmp_index[x], tmp_index[x + 1]):
                        tmp.loc[line, '理财产品/内部投组名称'] = title
                tmp = tmp[~tmp['投资资产分类'].isna()]
                tmp['市值'] = tmp['投资金额(万元)'] / 10000
                tmp['资产名称'].replace(float('nan'), "", inplace=True)
                tmp['资产名称'] = tmp['资产名称'].map(lambda lambda_x: lambda_x[:lambda_x.rfind("(")])
                tmp = tmp[['理财产品/内部投组名称', '投资资产分类', '资产名称', '市值', '到期日', '起息日', '买入价格', '百元估值']]
                tmp.columns = ['投组单元名称', '产品分类', '名称', '市值', '到期日', '起息日', '成本', '估值']
                tmp['业务日期'] = self.data.loc[0, '业务日期']
                for x in instrument:
                    tmp.loc[tmp['名称'] == x, '产品分类'] = '理财直融工具'
                self.data = self.data.append(tmp)
                self.data = self.data.reset_index(drop=True)

        def prepare_stream(self):
            if os.path.exists('data/现券交易.xls'):
                tmp = pd.read_excel("data/现券交易.xls", header=1)
                tmp = tmp[(tmp['Unnamed: 23'] == '成交确认') | (tmp['Unnamed: 23'] == '分配完毕')]
                tmp['名称'] = tmp['Unnamed: 14']
                tmp['类别'] = '债券'
                tmp['交易日'] = tmp['Unnamed: 11']
                tmp['方向'] = tmp['Unnamed: 3']
                tmp['金额'] = tmp['Unnamed: 20'] / 100000000
                tmp['交易投组'] = tmp['Unnamed: 4']
                tmp['对手方'] = tmp['对手名称']
                tmp['净价'] = tmp['Unnamed: 16']
                tmp.sort_values(by=['交易日'], inplace=True)
                tmp = tmp.reset_index(drop=True)
                db = pymysql.connect(host='localhost', port=3306, user='root', password='root', db='cost',
                                     charset='utf8')
                cost = pd.read_sql("select part,name,cost,amount from licai", db)
                cost['match'] = cost['part'] + cost['name']
                tmp['match'] = tmp['交易投组'] + tmp['名称']
                for x in range(len(tmp)):
                    if tmp.loc[x, '方向'] == '现券买入':
                        cost_match = cost.loc[cost['match'] == tmp.loc[x, 'match'], ['cost', 'amount']]
                        cost_match = cost_match.reset_index(drop=True)
                        if len(cost_match) > 0:
                            cost.loc[cost['match'] == tmp.loc[x, 'match'], 'amount'] = tmp.loc[x, 'Unnamed: 15'] / 100 \
                                                                                       + cost_match.loc[0, 'amount']
                            cost.loc[cost['match'] == tmp.loc[x, 'match'], 'cost'] = (tmp.loc[x, 'Unnamed: 15'] / 100 *
                                                                                      tmp.loc[x, '净价'] +
                                                                                      cost_match.loc[0, 'cost'] *
                                                                                      cost_match.loc[0, 'amount']) / \
                                                                                     (tmp.loc[x, 'Unnamed: 15'] / 100 +
                                                                                      cost_match.loc[0, 'amount'])
                        else:
                            cost.loc[cost['match'] == tmp.loc[x, 'match'], 'amount'] = tmp.loc[x, 'Unnamed: 15'] / 100
                            cost.loc[cost['match'] == tmp.loc[x, 'match'], 'cost'] = tmp.loc[x, '净价']
                    else:
                        res = cost.loc[cost['match'] == tmp.loc[x, 'match'], 'amount'].sum() - tmp.loc[
                            x, 'Unnamed: 15'].sum() / 100
                        if res <= 0:
                            cost.loc[cost['match'] == tmp.loc[x, 'match'], 'amount'] = 0
                            cost.loc[cost['match'] == tmp.loc[x, 'match'], 'cost'] = 0
                        else:
                            cost.loc[cost['match'] == tmp.loc[x, 'match'], 'amount'] = res
                cost = cost[cost['cost'] > 0]
                cost = cost.reset_index(drop=True)
                cost = cost[['part', 'name', 'cost', 'amount']]
                cur = db.cursor()
                cur.execute("delete from licai")
                for x in range(len(cost)):
                    cur.execute("insert into licai values('" + str(cost.loc[x, 'part']) + "','" + str(
                        cost.loc[x, 'name']) + "','" + str(cost.loc[x, 'cost']) + "','" + str(
                        cost.loc[x, 'amount']) + "')")
                db.commit()
                db.close()
                tmp = tmp[['名称', '类别', '交易日', '方向', '金额', '交易投组', '对手方', '净价']]
                self.flow = self.flow.append(tmp)
                self.flow = self.flow.reset_index(drop=True)

            if os.path.exists('data/债券交易.xls'):
                tmp = pd.read_excel("data/债券交易.xls")
                tmp['名称'] = tmp['债券简称']
                tmp['类别'] = '债券'
                tmp['交易日'] = tmp['交割日'].map(lambda lambda_x: pd.to_datetime(lambda_x, format='%Y-%m-%d'))
                tmp['方向'] = tmp['交易方向']
                tmp['金额'] = tmp['交易金额'] / 100000000
                tmp['交易投组'] = tmp['账户'].map(lambda lambda_x: lambda_x[:lambda_x.rfind("(")])
                tmp['对手方'] = tmp['交易对手']
                tmp['净价'] = (tmp['全价总额'] - tmp['应计利息总额']) / tmp['券面总额'] * 100
                tmp = tmp[['名称', '类别', '交易日', '方向', '金额', '交易投组', '对手方', '净价']]
                self.flow = self.flow.append(tmp)
                self.flow = self.flow.reset_index(drop=True)

            if os.path.exists('data/债券质押式回购交易.xls'):
                tmp = pd.read_excel("data/债券质押式回购交易.xls")
                tmp['名称'] = '回购'
                tmp['类别'] = '回购'
                tmp['交易日'] = tmp['起息日'].map(lambda lambda_x: pd.to_datetime(lambda_x, format='%Y-%m-%d'))
                tmp['方向'] = tmp['回购方向']
                tmp['金额'] = tmp['成交金额'] / 100000000
                tmp['交易投组'] = tmp['账户'].map(lambda lambda_x: lambda_x[:lambda_x.rfind("(")])
                tmp['对手方'] = tmp['交易对手']
                tmp = tmp[['名称', '类别', '交易日', '方向', '金额', '交易投组', '对手方']]
                self.flow = self.flow.append(tmp)
                self.flow = self.flow.reset_index(drop=True)

            if os.path.exists('data/质押式回购交易.xls'):
                tmp = pd.read_excel("data/质押式回购交易.xls", header=1)
                tmp = tmp[(tmp['Unnamed: 24'] == '成交确认') | (tmp['Unnamed: 24'] == '分配完毕')]
                tmp['名称'] = tmp['Unnamed: 12']
                tmp['类别'] = '回购'
                tmp['交易日'] = tmp['Unnamed: 11']
                tmp['方向'] = tmp['Unnamed: 3'].replace({'质押式逆回购': '逆回购', '质押式正回购': '正回购'})
                tmp['金额'] = tmp['结算金额(元)'] / 100000000
                tmp['交易投组'] = tmp['Unnamed: 4']
                tmp['对手方'] = tmp['对手名称']
                tmp = tmp[['名称', '类别', '交易日', '方向', '金额', '交易投组', '对手方']]
                self.flow = self.flow.append(tmp)
                self.flow = self.flow.reset_index(drop=True)

    def __init__(self, name):
        self.name = {"同业业务中心": self.TY, "理财事业部": self.LC}.get(name)()

    def get_asset(self):
        return self.name.asset

    def get_loan(self):
        return self.name.loan

    def get_stream(self):
        return self.name.stream


class BankInfo:
    def __init__(self, date_date):
        year = date_date.year
        month = (date_date.month - 1) // 3 * 3
        if month == 0:
            year -= 1
            month = 12
        day = {3: 31, 6: 30, 9: 30, 12: 31}.get(month)
        self.date = str(year) + "/" + str(month) + "/" + str(day)
        db = pymysql.connect(host='localhost', port=3306, user='root', password='root', db='pac', charset='utf8')
        data = pd.read_sql("select name,val/10000 from blancesheet where date='" + self.date + "'", db)
        self.asset = float(data.loc[data['name'] == '总资产', 'val/10000'].sum())
        self.loan = float(data.loc[data['name'] == '总负债', 'val/10000'].sum())
        self.net = self.asset - self.loan
        self.capital = float(data.loc[data['name'] == '一级资本净额', 'val/10000'].sum())

    def get_date(self):
        return self.date

    def get_asset(self):
        return self.asset

    def get_net(self):
        return self.net

    def get_capital(self):
        return self.capital

    def get_anonymous_special(self):
        db = pymysql.connect(host='localhost', port=3306, user='root', password='root', db='pac', charset='utf8')
        return pd.read_sql('select val from anonymous where date=' + self.date + "'", db).loc[0, 'val']


class Bond:
    def __init__(self, bond):
        self.bond = bond
        self.flag = False

    def code(self):
        db = pymysql.connect(host='localhost', port=3306, user='root', password='root', db='pac', charset='utf8')
        code = pd.read_sql("select * from bond_code where name in ('" + "','".join(self.bond['名称'].tolist()) + "')", db)
        data = pd.merge(self.bond, code, how="left", left_on="名称", right_on="name")
        cur = db.cursor()
        for line in data[data['code'].isna()].index.tolist():
            new_name = data.loc[line, '名称']
            new_code = input(new_name + "对应的代码：")
            cur.execute("insert into bond_code values('" + new_name + "','" + new_code + "')")
            data.loc[line, 'code'] = new_code
        db.commit()
        db.close()
        data['债券代码'] = data['code']
        self.bond = data[self.bond.columns.tolist() + ['债券代码']].copy(deep=True)

    def asset_bond(self):
        if self.flag:
            return self.bond
        self.code()
        if not w.isconnected():
            w.start()
        data = w.wss(",".join(set(self.bond['债券代码'].tolist())),
                     "windl1type,windl2type,province,city,comp_name,municipalbond,subordinateornot,mixcapital,perpetualornot,issue_issuemethod,modidura_cnbd,net_cnbd,latestpar,amount,latestissurercreditrating,issueamount,ptmyear",
                     "unit=1;tradeDate=" + str(self.bond.loc[0, '业务日期']).split()[0].replace("-", "") + ";credibility=1",
                     usedf=True)[1]
        wind_columns = ['WIND一级分类', 'WIND二级分类', '省份', '城市', '发行主体', '是否城投债', '是否次级债', '是否混合资本债券',
                        '是否永续债', '发行方式', '修正久期', '估值净价', '最新面值', '债项评级', '主体评级', '发行总额', '剩余期限']
        data.columns = wind_columns
        data = pd.merge(self.bond, data, how='left', left_on='债券代码', right_index=True)
        self.bond = data[self.bond.columns.tolist() + wind_columns]
        self.bond.loc[self.bond['WIND一级分类'].str.contains('国债|地方政府债|央行票据|政府支持机构债'), '债券类别'] = '利率债'
        self.bond.loc[(self.bond['WIND一级分类'] == '金融债') & (self.bond['WIND二级分类'] == '政策银行债'), '债券类别'] = '利率债'
        self.bond.loc[(self.bond['债券类别'].isna()) & (self.bond['WIND一级分类'] == '金融债'), '债券类别'] = '金融债'
        self.bond.loc[self.bond['WIND一级分类'] == '同业存单', '债券类别'] = '金融债'
        self.bond.loc[self.bond['债券类别'].isna(), '债券类别'] = '非金融企业债券'
        db = pymysql.connect(host='localhost', port=3306, user='root', password='root', db='pac', charset='utf8')
        special = pd.read_sql("select * from bond_special_info", db)
        for name in special['code'].tolist():
            self.bond.loc[self.bond['名称'] == name, '发行主体'] = special.loc[special['code'] == name, 'name']
            self.bond.loc[self.bond['名称'] == name, '是否城投债'] = special.loc[
                special['code'] == name, 'municipalbond']
            self.bond.loc[self.bond['名称'] == name, '发行方式'] = special.loc[
                special['code'] == name, 'issuemethod']
            self.bond.loc[self.bond['名称'] == name, 'WIND一级分类'] = special.loc[
                special['code'] == name, 'windl1type']
            self.bond.loc[self.bond['名称'] == name, '省份'] = special.loc[special['code'] == name, 'province']
            self.bond.loc[self.bond['名称'] == name, '城市'] = special.loc[special['code'] == name, 'city']
            self.bond.loc[self.bond['名称'] == name, '最新面值'] = special.loc[special['code'] == name, 'latestpar']
            self.bond.loc[self.bond['名称'] == name, '发行总额'] = special.loc[
                special['code'] == name, 'issueamount']
        if os.path.exists("data/指定成本与FIFO损益分析-新.xls"):
            res = pd.read_excel("data/指定成本与FIFO损益分析-新.xls")
            res = res[['Unnamed: 4', '市价修正久期']]
            for name in special['code'].tolist():
                self.bond.loc[self.bond['名称'] == name, '修正久期'] = res.loc[res['Unnamed: 4'] == name, '市价修正久期']
        self.flag = True
        return self.bond

    def asset_credit(self):
        if not self.flag:
            self.asset_bond()
        return self.bond[self.bond['债券类别'] == '非金融企业债券']

    def asset_abs(self):
        if not self.flag:
            self.asset_bond()
        data = self.bond[self.bond['WIND一级分类'] == '资产支持证券']
        if not w.isconnected():
            w.start()
        abs_type = w.wss(",".join(set(data['债券代码'].tolist())), "us_type", usedf=True)[1]
        abs_type.columns = ['ABS基础资产类型']
        data = pd.merge(data, abs_type, how='left', left_on='债券代码', right_index=True)
        return data

    def stream_get(self):
        if self.flag:
            return self.bond
        self.code()
        if not w.isconnected():
            w.start()
        for x in self.bond.index.tolist():
            self.bond.loc[x, '估值净价'] = w.wss(self.bond.loc[x, '债券代码'], "net_cnbd", "tradeDate=" +
                                             str(self.bond.loc[x, '交易日']).split()[0].replace("-", "") +
                                             ";credibility=1").Data[0][0]
        self.flag = True
        return self.bond


class MMF:
    def __init__(self, data):
        self.data = data

    def wind(self):
        db = pymysql.connect(host='localhost', port=3306, user='root', password='root', db='pac', charset='utf8')
        code = pd.read_sql("select name,code from mmf_code where name in ('" + "','".join(self.data['名称'].tolist()) +
                           "')", db)
        if not w.isconnected():
            w.start()
        data = w.wss(",".join(code['code'].tolist()), "fund_corp_fundmanagementcompany,fund_fundscale", "unit=1",
                     usedf=True)[1]
        data.columns = ['管理人', '基金规模']
        self.data = pd.merge(self.data, code, left_on='名称', right_on='name')
        self.data = pd.merge(self.data, data, left_on='code', right_index=True)
        self.data['占比'] = self.data['市值'] * 100000000 / self.data['基金规模']

    def ratio(self):
        self.wind()
        out = [str(round(self.data['占比'].max() * 100, 2)),
               str(round(self.data['占比'].min() * 100, 2)),
               str(len(set(self.data['管理人'].tolist()))) + "家"]
        res = self.data.groupby('管理人')['市值'].sum()
        out.append(str(round(res.max(), 2)))
        out.append(str(round(res.min(), 2)))
        return out


class ETF:
    def __init__(self, data):
        self.etf = data.bs.asset[data.bs.asset['名称'].str.contains('etf|ETF')]
        self.net = data.bs.loan[data.bs.loan['产品分类'] == '理财产品'].groupby('投组单元名称', as_index=False)['市值'].sum()
        self.net = self.net[['投组单元名称', '净资产']]
        self.date = str(data.bs.asset.loc[0, '业务日期']).replace("-", "")
        self.flag = False

    def wind(self):
        if not self.flag:
            etf = list(set(self.etf['名称'].tolist()))
            db = pymysql.connect(host='localhost', port=3306, user='root', password='root', db='pac', charset='utf8')
            cur = db.cursor()
            code = pd.read_sql("select name,code from partner_code where name in ('" + "','".join(etf) + "')", db)
            for x in etf:
                if x not in code['name'].tolist():
                    new_code = input("请补充" + x + "对应的代码：")
                    code.loc[x, 'code'] = new_code
                    cur.execute("insert into partner_code values('" + x + "','" + new_code + "'")
            db.commit()
            if not w.isconnected():
                w.start()
            etf = w.wss(",".join(code['code'].tolist()), "close", "tradeDate=" + self.date + ";priceAdj=U;cycle=D",
                        usedf=True)
            self.etf = pd.merge(self.etf, code, how='left', left_on='名称', right_on='name')
            self.etf = pd.merge(self.etf, etf, how='left', left_on='code', right_index=True)
            self.flag = True

    def get(self):
        if not self.flag:
            self.wind()
        ratio = self.etf.groupby('投组单元名称', as_index=False)['市值'].sum()
        ratio.columns = ['投组单元名称', 'etf']
        ratio = pd.merge(ratio, self.net, how='left', on='投组单元名称')
        ratio['ratio'] = ratio['etf'] / ratio['净资产']
        db = pymysql.connect(host='localhost', port=3306, user='root', password='root', db='cost', charset='utf8')
        cost = pd.read_sql("select name,cost from cost_etf", db)
        self.etf = pd.merge(self.etf, cost, how='left', left_on='名称', right_on='name')
        earn = ((self.etf['close'] - self.etf['cost']) / self.etf['cost']).max() * 100
        lost = ((self.etf['close'] - self.etf['cost']) / self.etf['cost']).min() * 100
        out = [
            str(round(self.etf['市值'].sum(), 2)),
            "≤" + str(round(self.etf['市值'].max(), 2)),
            "≤" + str(round(ratio['ratio'].max(), 2)),
            [round(earn, 2), round(lost, 2)]
        ]
        return out


class Party:
    def __init__(self, ty, lc, fix_date):
        self.fix_date = str(pd.to_datetime(fix_date, format='%Y/%m%d')).replace('-', '')
        ty['方向'] = ty['类别'] + ty['方向']
        ty['部门'] = '自营'
        self.data = ty[['对手方', '方向', '部门', '金额']].copy(deep=True)
        lc['方向'] = ty['类别'] + ty['方向']
        lc['部门'] = '理财'
        lc = lc[['对手方', '方向', '部门', '金额']].copy(deep=True)
        self.data = self.data.append(lc)

    def wind(self):
        partner = list(set(self.data['对手方'].tolist()))
        db = pymysql.connect(host='localhost', port=3306, user='root', password='root', db='pac', charset='utf8')
        cur = db.cursor()
        code = pd.read_sql("select name,code from partner_code where name in ('" + "','".join(partner) + "')", db)
        for x in partner:
            if x not in code['name'].tolist():
                new_code = input("请补充" + x + "对应的代码：")
                code.loc[x, 'code'] = new_code
                cur.execute("insert into partner_code values('" + x + "','" + new_code + "'")
        db.commit()
        bad = pd.merge(code[code['code'] == 'BAD'], self.data, how='left', left_on='名称', right_on='name')
        bad = bad[['对手方', '方向', '部门', '金额']]
        code = code[(code['code'] != 'BAD') & (code['code'] != 'GOOD')]
        if not w.isconnected():
            w.start()
        tmp = w.wss(",".join(code['code'].tolist()), "regcapital,latestissurercreditrating",
                    "tradeDate=" + self.fix_date + ";industryType=3;unit=1", usedf=True)[1]
        tmp = tmp[(tmp['REGCAPITAL'] < 100000000000) & (tmp['LATESTISSURERCREDITRATING'] != 'AAA')]
        if len(tmp) > 0:
            self.data = pd.merge(code, self.data, how='left', left_on='name', right_on='名称')
            self.data = pd.merge(tmp, self.data, how='left', left_on='code', right_index=True)
            self.data = self.data[['对手方', '方向', '部门', '金额']]
            self.data = self.data.append(bad)
        elif len(bad) > 0:
            self.data = bad
        else:
            self.data = pd.DataFrame()

    def get(self):
        out = None
        if len(self.data) > 0:
            out = self.data.groupby(['对手方', '方向', '部门'])['金额'].agg(['sum', 'count', 'max', 'min'])
        return out


class BalanceSheet:
    def __init__(self, asset, loan):
        self.asset = asset()
        if '业务日期' not in self.asset.columns.tolist():
            self.asset['业务日期'] = pd.to_datetime(input("请输入资产日期（%Y-%m-%d）："), format='%Y-%m-%d')
        self.loan = loan()
        self.bond = Bond(self.asset[self.asset['产品分类'] == '债券'].copy(deep=True))


class Stream:
    def __init__(self, stream):
        self.stream = stream()
        self.bond = Bond(self.stream[self.stream['类别'] == '债券'].copy(deep=True)).stream_get()


class Department:
    def __init__(self, name):
        self.name = name
        func = Function(name)
        self.stream = Stream(func.get_stream())
        self.bs = BalanceSheet(func.get_asset(), func.get_loan())

    def struct(self):
        data = {}
        tmp = self.bs.asset.groupby('产品分类', as_index=False)['市值'].sum()
        tmp['占比'] = tmp['市值'] / tmp['市值'].sum()
        data['资产'] = tmp
        tmp = self.bs.loan.groupby('产品分类', as_index=False)['市值'].sum()
        tmp['占比'] = tmp['市值'] / tmp['市值'].sum()
        data['负债'] = tmp
        return data

    def concentration(self):
        res = self.bs.bond.asset_credit()
        out = []
        data = res.groupby('发行主体')['市值'].sum()
        big = []
        tmp = list(set(data.sort_values(ascending=False).index.tolist()))[:10]
        for x in tmp:
            big.append([x, str(round(data[x], 2)), str(len(res[res['发行主体'] == x]))])
        level = [set(data[data <= 0.5].index.tolist()), set(data[(data > 0.5) & (data <= 1)].index.tolist()),
                 set(data[(data > 1) & (data <= 2)].index.tolist()), set(data[data > 2].index.tolist())]
        for x in range(len(level)):
            tmp = res[res['发行主体'].isin(level[x])]
            out.append([str(len(level[x])), str(len(tmp)), str(round(tmp['市值'].sum(), 2)),
                        str(round(tmp['市值'].sum() / res['市值'].sum() * 100, 2)) + '%'])
        return out, big

    def area(self):
        res = self.bs.bond.asset_credit()
        res = res[res['省份'] == '福建省']
        data = res.groupby('城市', as_index=False)['市值'].sum()
        data['占比'] = (data['市值'] / data['市值'].sum() * 100).map(lambda x: str(round(x, 2)) + '%')
        data['市值'] = data['市值'].map(lambda x: str(round(x, 2)))
        return data

    def duration(self):
        res = self.bs.bond.asset_bond()
        return [[round(res['市值'].sum(), 2), round((res['市值'] * res['修正久期']).sum() / res['市值'].sum(), 2)],
                [round(res[res['债券类别'] == '利率债']['市值'].sum(), 2),
                 round((res[res['债券类别'] == '利率债']['市值'] * res['修正久期']).sum() /
                       res[res['债券类别'] == '利率债']['市值'].sum(), 2)],
                [round(res[res['债券类别'] == '金融债']['市值'].sum(), 2),
                 round((res[res['债券类别'] == '金融债']['市值'] * res['修正久期']).sum() /
                       res[res['债券类别'] == '金融债']['市值'].sum(), 2)],
                [round(res[res['债券类别'] == '非金融企业债券']['市值'].sum(), 2),
                 round((res[res['债券类别'] == '非金融企业债券']['市值'] * res['修正久期']).sum() /
                       res[res['债券类别'] == '非金融企业债券']['市值'].sum(), 2)]]

    def ratio(self, total):
        out = []
        res = self.bs.bond.asset_bond()
        res['评级'] = res['主体评级']
        res.loc[res['评级'].isna(), '评级'] = res.loc[res['评级'].isna(), '主体评级']
        res['评级AA+以下'] = '是'
        res.loc[res['评级'] == 'AAA', '评级AA+以下'] = '否'
        res.loc[res['评级'] == 'AA+', '评级AA+以下'] = '否'
        total_bond = res['市值'].sum()
        out.append(str(round(total_bond / total * 100, 2)) + "%")
        total_credit = res.loc[res['债券类别'] == '非金融企业债券', '市值'].sum()
        out.append(str(round(total_credit / total_bond * 100, 2)) + "%")
        out.append(str(round(res.loc[res['评级AA+以下'] == '是', '市值'].sum() / total_bond * 100, 2)) + "%")
        out.append(str(round(res.loc[res['是否城投债'] == '是', '市值'].sum() / total_bond * 100, 2)) + "%")
        out.append(str(round(res.loc[res['发行方式'] == '私募', '市值'].sum() / total_bond * 100, 2)) + "%")
        out.append(str(round(res.loc[res['是否永续债'] == '是', '市值'].sum() / total_bond * 100, 2)) + "%")
        out.append(str(round(res.loc[res['是否次级债'] == '是', '市值'].sum() / total_bond * 100, 2)) + "%")
        out.append(str(round(res.loc[res['是否混合资本债券'] == '是', '市值'].sum() / total_bond * 100, 2)) + "%")
        return out

    def lost(self):
        res = self.bs.bond.asset_bond()
        res['成本'] = res['成本'] / res['最新面值'] * 100
        res['止损限额'] = (res['成本'] - res['估值净价']) / res['成本'] * 100
        return [res.loc[(res['剩余期限'] <= 5) & (res['止损限额'] > 16), :],
                res.loc[(res['剩余期限'] > 5) & (res['剩余期限'] <= 10) & (res['止损限额'] > 23), :],
                res.loc[(res['剩余期限'] > 10) & (res['止损限额'] > 40), :]]

    def credit_limit(self):
        res = self.bs.bond.asset_credit()
        res = res[res['WIND一级分类'] != '资产支持证券']
        res = pd.merge(res.groupby('名称', as_index=False)['市值'].sum(), res[['名称', '发行总额']], how='left', on='名称')
        res['单券集中度'] = res['市值'] / res['发行总额'] * 100000000
        return [res['单券集中度'].max(), res[res['单券集中度'] > 0.2]]

    def deviate(self):
        res = self.stream.bond
        res['交易价格偏离度'] = abs((res['净价'] - res['估值净价']) / res['估值净价'])
        return res['交易价格偏离度'].max()

    def stream_description(self):
        res = self.stream.stream
        return [res.groupby(['类别', '方向'])['金额'].agg(['count', 'sum', 'max', 'min']),
                res.groupby(['交易投组', '方向'])['金额'].agg(['count', 'sum', 'max', 'min'])]


class Word:
    def __init__(self):
        self.document = Document('风险管理部金融市场风险监测报告模板.docx')

    @staticmethod
    def delete_row(table, n):
        row = table.rows[n]
        tbl = table._tbl
        tr = row._tr
        tbl.remove(tr)

    def sharp_table(self, table, n):
        total = len(table.rows)
        if total > n:
            for x in range(total - n):
                self.delete_row(table, 2)
        elif total < n:
            raise ValueError

    @staticmethod
    def style_cell(cell, name, size, bold=False):
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            for run in paragraph.runs:
                run.font.name = name
                run.font.size = size
                run.font.bold = bold

    def leverage_ty(self, data, info):
        self.document.tables[20].cell(2, 1).text = str(
            round(data.bs.loan.loc[data.bs.loan['产品分类'] == '回购', '市值'].sum() / info.get_net(), 2))
        self.style_cell(self.document.tables[20].cell(2, 1), '宋体', 177800)
        self.document.tables[20].cell(3, 1).text = str(
            round(data.bs.asset.loc[data.bs.asset['产品分类'] == '回购', '市值'].sum() / info.get_net(), 2))
        self.style_cell(self.document.tables[20].cell(3, 1), '宋体', 177800)

    def leverage_lc(self, data):
        name_list = list(set(data.bs.asset['投组单元名称'].tolist()))
        for x in range(len(name_list)):
            name = name_list[x]
            total = data.bs.asset.loc[data.bs.asset['投组单元名称'] == name, '市值'].sum()
            self.document.tables[20].add_row()
            self.document.tables[20].cell(5 + 3 * x, 0).text = name + "杠杆率"
            self.style_cell(self.document.tables[20].cell(5 + 3 * x, 0), '宋体', 177800)
            self.document.tables[20].cell(5 + 3 * x, 1).text = str(round(total / data.bs.loan.loc[
                (data.bs.loan['投组单元名称'] == name) & (data.bs.loan['产品分类'] == '理财产品'), '市值'].sum(), 2))
            self.style_cell(self.document.tables[20].cell(5 + 3 * x, 1), '宋体', 177800)
            self.document.tables[20].add_row()
            self.document.tables[20].cell(6 + 3 * x, 0).text = name + "逆回购杠杆率"
            self.style_cell(self.document.tables[20].cell(6 + 3 * x, 0), '宋体', 177800)
            self.document.tables[20].cell(6 + 3 * x, 1).text = str(
                round(data.bs.asset.loc[(data.bs.asset['投组单元名称'] == name) & (
                        data.bs.asset['产品分类'] == '买入返售金融资产'), '市值'].sum() / total, 2))
            self.style_cell(self.document.tables[20].cell(6 + 3 * x, 1), '宋体', 177800)
            self.document.tables[20].add_row()
            self.document.tables[20].cell(7 + 3 * x, 0).text = name + "正回购杠杆率"
            self.style_cell(self.document.tables[20].cell(7 + 3 * x, 0), '宋体', 177800)
            self.document.tables[20].cell(7 + 3 * x, 1).text = str(
                round(data.bs.loan.loc[(data.bs.loan['投组单元名称'] == name) & (
                        data.bs.loan['产品分类'] == '卖出回购金融资产款'), '市值'].sum() / total, 2))
            self.style_cell(self.document.tables[20].cell(7 + 3 * x, 1), '宋体', 177800)
        self.delete_row(self.document.tables[20], len(self.document.tables[20].rows) - 1)

    def go(self):
        ty = Department("同业业务中心")
        lc = Department("理财事业部")
        bank = BankInfo(ty.bs.asset.loc[0, '业务日期'].date())

        size = len(ty.struct().get('资产'))
        if len(ty.struct().get('负债')) > size:
            size = len(ty.struct().get('负债'))
        self.sharp_table(self.document.tables[4], size + 4)
        for x in range(len(ty.struct().get('资产'))):
            self.document.tables[4].cell(2 + x, 0).text = ty.struct().get('资产').loc[x, '产品分类']
            self.style_cell(self.document.tables[4].cell(2 + x, 0), '宋体', 177800)
            self.document.tables[4].cell(2 + x, 1).text = str(round(ty.struct().get('资产').loc[x, '市值'], 2))
            self.style_cell(self.document.tables[4].cell(2 + x, 1), '宋体', 177800)
            self.document.tables[4].cell(2 + x, 2).text = str(round(ty.struct().get('资产').loc[x, '占比'] * 100, 2)) + "%"
            self.style_cell(self.document.tables[4].cell(2 + x, 2), '宋体', 177800)
        self.document.tables[4].cell(2 + size, 1).text = str(round(ty.struct().get('资产')['市值'].sum(), 2))
        self.style_cell(self.document.tables[4].cell(2 + size, 1), '宋体', 177800, True)
        for x in range(len(ty.struct().get('负债'))):
            self.document.tables[4].cell(2 + x, 3).text = ty.struct().get('负债').loc[x, '产品分类']
            self.style_cell(self.document.tables[4].cell(2 + x, 3), '宋体', 177800)
            self.document.tables[4].cell(2 + x, 4).text = str(round(ty.struct().get('负债').loc[x, '市值'], 2))
            self.style_cell(self.document.tables[4].cell(2 + x, 4), '宋体', 177800)
            self.document.tables[4].cell(2 + x, 5).text = str(round(ty.struct().get('负债').loc[x, '占比'] * 100, 2)) + "%"
            self.style_cell(self.document.tables[4].cell(2 + x, 5), '宋体', 177800)
        self.document.tables[4].cell(2 + size, 4).text = str(round(ty.struct().get('负债')['市值'].sum(), 2))
        self.style_cell(self.document.tables[4].cell(2 + size, 4), '宋体', 177800, True)

        size = len(lc.struct().get('资产'))
        if len(lc.struct().get('负债')) > size:
            size = len(lc.struct().get('负债'))
        self.sharp_table(self.document.tables[5], size + 4)
        for x in range(len(lc.struct().get('资产'))):
            self.document.tables[5].cell(2 + x, 0).text = lc.struct().get('资产').loc[x, '产品分类']
            self.style_cell(self.document.tables[5].cell(2 + x, 0), '宋体', 177800)
            self.document.tables[5].cell(2 + x, 1).text = str(round(lc.struct().get('资产').loc[x, '市值'], 2))
            self.style_cell(self.document.tables[5].cell(2 + x, 1), '宋体', 177800)
            self.document.tables[5].cell(2 + x, 2).text = str(round(lc.struct().get('资产').loc[x, '占比'] * 100, 2)) + "%"
            self.style_cell(self.document.tables[5].cell(2 + x, 2), '宋体', 177800)
        self.document.tables[5].cell(2 + size, 1).text = str(round(lc.struct().get('资产')['市值'].sum(), 2))
        self.style_cell(self.document.tables[5].cell(2 + size, 1), '宋体', 177800, True)
        for x in range(len(lc.struct().get('负债'))):
            self.document.tables[5].cell(2 + x, 3).text = lc.struct().get('负债').loc[x, '产品分类']
            self.style_cell(self.document.tables[5].cell(2 + x, 3), '宋体', 177800)
            self.document.tables[5].cell(2 + x, 4).text = str(round(lc.struct().get('负债').loc[x, '市值'], 2))
            self.style_cell(self.document.tables[5].cell(2 + x, 4), '宋体', 177800)
            self.document.tables[5].cell(2 + x, 5).text = str(round(lc.struct().get('负债').loc[x, '占比'] * 100, 2)) + "%"
            self.style_cell(self.document.tables[5].cell(2 + x, 5), '宋体', 177800)
        self.document.tables[5].cell(2 + size, 4).text = str(round(lc.struct().get('负债')['市值'].sum(), 2))
        self.style_cell(self.document.tables[5].cell(2 + size, 4), '宋体', 177800, True)

        data, big = ty.concentration()
        self.document.tables[1].cell(1, 1).text = str(sum([int(x[0]) for x in data]))
        self.style_cell(self.document.tables[0].cell(1, 1), '宋体', 177800)
        self.document.tables[13].cell(6, 1).text = str(sum([int(x[0]) for x in data]))
        self.style_cell(self.document.tables[0].cell(1, 1), '宋体', 177800, True)
        for x in range(4):
            for y in range(4):
                self.document.tables[13].cell(2 + x, 1 + y).text = data[x][y]
                self.style_cell(self.document.tables[13].cell(2 + x, 1 + y), '宋体', 177800)
        for x in range(10):
            for y in range(3):
                self.document.tables[14].cell(2 + x, y).text = big[x][y]
                self.style_cell(self.document.tables[14].cell(2 + x, y), '宋体', 177800)
        data, big = lc.concentration()
        for x in range(4):
            for y in range(4):
                self.document.tables[16].cell(2 + x, 1 + y).text = data[x][y]
                self.style_cell(self.document.tables[16].cell(2 + x, 1 + y), '宋体', 177800)
        for x in range(10):
            for y in range(3):
                self.document.tables[17].cell(2 + x, y).text = big[x][y]
                self.style_cell(self.document.tables[17].cell(2 + x, y), '宋体', 177800)

        data = ty.area()
        for x in range(len(data)):
            for y in range(3):
                self.document.tables[15].cell(2 + x, y).text = str(data.loc[x, data.columns.tolist()[y]])
                self.style_cell(self.document.tables[15].cell(2 + x, y), '宋体', 177800)
        data = lc.area()
        for x in range(len(data)):
            for y in range(3):
                self.document.tables[18].cell(2 + x, y).text = str(data.loc[x, data.columns.tolist()[y]])
                self.style_cell(self.document.tables[18].cell(2 + x, y), '宋体', 177800)

        data = ty.duration()
        for x in range(4):
            for y in range(2):
                self.document.tables[19].cell(1 + x, 1 + y).text = str(data[x][y])
                self.style_cell(self.document.tables[19].cell(1 + x, 1 + y), '宋体', 177800)
        data = lc.duration()
        for x in range(4):
            for y in range(2):
                self.document.tables[19].cell(5 + x, 1 + y).text = str(data[x][y])
                self.style_cell(self.document.tables[19].cell(5 + x, 1 + y), '宋体', 177800)

        self.leverage_ty(ty, bank)
        self.leverage_lc(lc)

        data = ty.ratio(bank.get_asset())
        for x in range(8):
            self.document.tables[21].cell(1 + x, 1).text = str(data[x])
            self.style_cell(self.document.tables[21].cell(1 + x, 1), '宋体', 177800)
        data = lc.ratio(bank.get_asset())
        for x in range(8):
            self.document.tables[21].cell(1 + x, 2).text = str(data[x])
            self.style_cell(self.document.tables[21].cell(1 + x, 2), '宋体', 177800)

        if len(ty.bs.asset[ty.bs.asset['产品分类'] == '货币基金']) > 0:
            data = MMF(ty.bs.asset[ty.bs.asset['产品分类'] == '货币基金']).ratio()
            for x in range(5):
                self.document.tables[22].cell(2 + x, 1).text = str(data[x])
                self.style_cell(self.document.tables[22].cell(2 + x, 1), '宋体', 177800)
        else:
            self.document.tables[22].cell(2, 1).text = "无业务"
            self.style_cell(self.document.tables[22].cell(2, 1), '宋体', 177800)

        self.document.tables[23].cell(1, 1).text = self.document.tables[19].cell(1, 2).text
        self.style_cell(self.document.tables[23].cell(1, 1), '宋体', 177800)
        self.document.tables[23].cell(2, 1).text = self.document.tables[19].cell(5, 2).text
        self.style_cell(self.document.tables[23].cell(2, 1), '宋体', 177800)
        data1 = float(self.document.tables[20].cell(2, 1).text[:-1])
        data2 = float(self.document.tables[20].cell(3, 1).text[:-1])
        if data1 > data2:
            self.document.tables[23].cell(3, 1).text = self.document.tables[20].cell(2, 1).text
        else:
            self.document.tables[23].cell(3, 1).text = self.document.tables[20].cell(3, 1).text
        self.style_cell(self.document.tables[23].cell(3, 1), '宋体', 177800)
        data1 = ty.deviate()
        data2 = lc.deviate()
        if data1 > data2:
            self.document.tables[23].cell(4, 1).text = '≤' + str(round(data1 * 100, 2)) + '%'
            self.style_cell(self.document.tables[23].cell(4, 1), '宋体', 177800)
        else:
            self.document.tables[23].cell(4, 1).text = '≤' + str(round(data2 * 100, 2)) + '%'
            self.style_cell(self.document.tables[23].cell(4, 1), '宋体', 177800)

        data_ty = ty.lost()
        data_lc = lc.lost()
        size = len(data_ty[0]) + len(data_ty[1]) + len(data_ty[2]) + len(data_lc[0]) + len(data_lc[1]) + len(data_lc[2])
        self.sharp_table(self.document.tables[24], size + 2)
        i = 1
        for x in data_ty:
            if len(x) > 0:
                for y in x.index.tolist():
                    self.document.tables[24].cell(i, 0).text = x.loc[y, '名称']
                    self.style_cell(self.document.tables[24].cell(i, 0), '宋体', 177800)
                    self.document.tables[24].cell(i, 1).text = '同业业务中心'
                    self.style_cell(self.document.tables[24].cell(i, 1), '宋体', 177800)
                    self.document.tables[24].cell(i, 2).text = str(x.loc[y, '止损限额']) + '%'
                    self.style_cell(self.document.tables[24].cell(i, 2), '宋体', 177800)
                    self.document.tables[24].cell(i, 3).text = ['≤16%', '≤23%', '≤40%'][x]
                    self.style_cell(self.document.tables[24].cell(i, 3), '宋体', 177800)
                    i += 1
        for x in data_lc:
            if len(x) > 0:
                for y in x.index.tolist():
                    self.document.tables[24].cell(i, 0).text = x.loc[y, '名称']
                    self.style_cell(self.document.tables[24].cell(i, 0), '宋体', 177800)
                    self.document.tables[24].cell(i, 1).text = '同业业务中心'
                    self.style_cell(self.document.tables[24].cell(i, 1), '宋体', 177800)
                    self.document.tables[24].cell(i, 2).text = str(x.loc[y, '止损限额']) + '%'
                    self.style_cell(self.document.tables[24].cell(i, 2), '宋体', 177800)
                    self.document.tables[24].cell(i, 3).text = ['≤16%', '≤23%', '≤40%'][x]
                    self.style_cell(self.document.tables[24].cell(i, 3), '宋体', 177800)
                    i += 1

        data_ty = ty.credit_limit()
        data_lc = lc.credit_limit()
        if data_ty[0] > 0.2:
            if data_lc[0] > 0.2:
                self.document.tables[25].cell(1, 1).text = str(data_ty[1]) + '\n' + str(data_lc[1])
            else:
                self.document.tables[25].cell(1, 1).text = str(data_ty[1])
        else:
            if data_lc[0] > 0.2:
                self.document.tables[25].cell(1, 1).text = str(data_lc[1])
            else:
                if data_ty[0] > data_lc[0]:
                    self.document.tables[25].cell(1, 1).text = '≤' + str(round(data_ty[0] * 100, 0))
                else:
                    self.document.tables[25].cell(1, 1).text = '≤' + str(round(data_lc[0] * 100, 0))
        self.style_cell(self.document.tables[25].cell(1, 1), '宋体', 177800)

        data = ty.bs.bond.asset_abs()
        res = [bank.get_anonymous_special(),
               ty.bs.asset.loc[ty.bs.asset['产品分类'] == '货币基金', '市值'].sum(),
               data.loc[(data['ABS基础资产类型'] == '航空票款') | (data['ABS基础资产类型'] == '基础设施收费'), '市值'].sum()]
        self.document.tables[26].cell(2, 2).text = str(round(res[0], 2))
        self.style_cell(self.document.tables[26].cell(2, 2), '宋体', 177800)
        self.document.tables[26].cell(3, 2).text = str(
            round(res[1], 2))
        self.style_cell(self.document.tables[26].cell(3, 2), '宋体', 177800)
        self.document.tables[26].cell(4, 2).text = str(round(res[2], 2))
        self.style_cell(self.document.tables[26].cell(4, 2), '宋体', 177800)
        res.append(sum(res))
        self.document.tables[26].cell(5, 1).text = str(round(res[3], 2))
        self.style_cell(self.document.tables[26].cell(5, 1), '宋体', 177800)
        res.append(bank.get_capital())
        self.document.tables[26].cell(6, 1).text = str(round(res[4], 2))
        self.style_cell(self.document.tables[26].cell(6, 1), '宋体', 177800)
        self.document.tables[26].cell(7, 1).text = str(round(res[3] / res[4] * 100, 2)) + '%'
        self.style_cell(self.document.tables[26].cell(7, 1), '宋体', 177800)

        data = ETF(lc).get()
        self.document.tables[27].cell(1, 1).text = data[0]
        self.style_cell(self.document.tables[27].cell(1, 1), '宋体', 177800)
        self.document.tables[27].cell(2, 1).text = data[1]
        self.style_cell(self.document.tables[27].cell(2, 1), '宋体', 177800)
        self.document.tables[27].cell(3, 1).text = data[2]
        self.style_cell(self.document.tables[27].cell(3, 1), '宋体', 177800)
        if data[3][0] < 0:
            self.document.tables[27].cell(4, 1).text = '无浮盈ETF'
            self.style_cell(self.document.tables[27].cell(4, 1), '宋体', 177800)
        else:
            self.document.tables[27].cell(4, 1).text = str(data[3][0])
            self.style_cell(self.document.tables[27].cell(4, 1), '宋体', 177800)
        if data[3][1] > 0:
            self.document.tables[27].cell(5, 1).text = '无浮亏ETF'
            self.style_cell(self.document.tables[27].cell(4, 1), '宋体', 177800)
        else:
            self.document.tables[27].cell(5, 1).text = str(-data[3][1])
            self.style_cell(self.document.tables[27].cell(4, 1), '宋体', 177800)

        data = ty.stream_description()
        self.sharp_table(self.document.tables[6], len(data[0]) + 2)
        self.sharp_table(self.document.tables[7], len(data[1]) + 2)
        self.document.tables[0].cell(1, 1).text = "、".join([x[0] for x in list(data[0].index)])
        self.style_cell(self.document.tables[0].cell(1, 1), '宋体', 177800)
        self.document.tables[0].cell(1, 2).text = str(round(data[0]['count'].sum(), 2))
        self.style_cell(self.document.tables[0].cell(1, 2), '宋体', 177800)
        self.document.tables[6].cell(len(data[0]) + 1, 1).text = str(round(data[0]['count'].sum(), 2))
        self.style_cell(self.document.tables[6].cell(len(data[0]) + 1, 1), '宋体', 177800, True)
        self.document.tables[0].cell(1, 3).text = str(round(data[0]['sum'].sum(), 2))
        self.style_cell(self.document.tables[0].cell(1, 3), '宋体', 177800)
        self.document.tables[6].cell(len(data[0]) + 1, 2).text = str(round(data[0]['sum'].sum(), 2))
        self.style_cell(self.document.tables[6].cell(len(data[0]) + 1, 2), '宋体', 177800, True)
        self.document.tables[0].cell(1, 4).text = str(round(data[0]['max'].max(), 2))
        self.style_cell(self.document.tables[0].cell(1, 4), '宋体', 177800)
        self.document.tables[6].cell(len(data[0]) + 1, 3).text = str(round(data[0]['max'].max(), 2))
        self.style_cell(self.document.tables[6].cell(len(data[0]) + 1, 3), '宋体', 177800, True)
        self.document.tables[0].cell(1, 5).text = str(round(data[0]['min'].min(), 2))
        self.style_cell(self.document.tables[0].cell(1, 5), '宋体', 177800)
        self.document.tables[6].cell(len(data[0]) + 1, 3).text = str(round(data[0]['min'].min(), 2))
        self.style_cell(self.document.tables[6].cell(len(data[0]) + 1, 3), '宋体', 177800, True)
        for x in range(len(data[0])):
            self.document.tables[6].cell(1 + x, 0).text = data[0].index[x][0]
            self.style_cell(self.document.tables[6].cell(1 + x, 0), '宋体', 177800)
            self.document.tables[6].cell(1 + x, 1).text = data[0].index[x][1]
            self.style_cell(self.document.tables[6].cell(1 + x, 1), '宋体', 177800)
            self.document.tables[6].cell(1 + x, 2).text = data[0].loc[data[0].index[x], 'count']
            self.style_cell(self.document.tables[6].cell(1 + x, 2), '宋体', 177800)
            self.document.tables[6].cell(1 + x, 3).text = data[0].loc[data[0].index[x], 'sum']
            self.style_cell(self.document.tables[6].cell(1 + x, 3), '宋体', 177800)
            self.document.tables[6].cell(1 + x, 4).text = data[0].loc[data[0].index[x], 'max']
            self.style_cell(self.document.tables[6].cell(1 + x, 4), '宋体', 177800)
            self.document.tables[6].cell(1 + x, 5).text = data[0].loc[data[0].index[x], 'min']
            self.style_cell(self.document.tables[6].cell(1 + x, 5), '宋体', 177800)
        for x in range(len(data[1])):
            self.document.tables[7].cell(1 + x, 0).text = data[1].index[x][0]
            self.style_cell(self.document.tables[7].cell(1 + x, 0), '宋体', 177800)
            self.document.tables[7].cell(1 + x, 1).text = data[1].index[x][1]
            self.style_cell(self.document.tables[7].cell(1 + x, 1), '宋体', 177800)
            self.document.tables[7].cell(1 + x, 2).text = data[1].loc[data[1].index[x], 'count']
            self.style_cell(self.document.tables[7].cell(1 + x, 2), '宋体', 177800)
            self.document.tables[7].cell(1 + x, 3).text = data[1].loc[data[1].index[x], 'sum']
            self.style_cell(self.document.tables[7].cell(1 + x, 3), '宋体', 177800)
            self.document.tables[7].cell(1 + x, 4).text = data[1].loc[data[1].index[x], 'max']
            self.style_cell(self.document.tables[7].cell(1 + x, 4), '宋体', 177800)
            self.document.tables[7].cell(1 + x, 5).text = data[1].loc[data[1].index[x], 'min']
            self.style_cell(self.document.tables[7].cell(1 + x, 5), '宋体', 177800)
        data = lc.stream_description()
        self.sharp_table(self.document.tables[8], len(data[0]) + 2)
        self.sharp_table(self.document.tables[9], len(data[1]) + 2)
        self.document.tables[0].cell(2, 1).text = "、".join([x[0] for x in list(data[0].index)])
        self.style_cell(self.document.tables[0].cell(2, 1), '宋体', 177800)
        self.document.tables[0].cell(2, 2).text = str(round(data[0]['count'].sum(), 2))
        self.style_cell(self.document.tables[0].cell(2, 2), '宋体', 177800)
        self.document.tables[7].cell(len(data[0]) + 1, 1).text = str(round(data[0]['count'].sum(), 2))
        self.style_cell(self.document.tables[7].cell(len(data[0]) + 1, 1), '宋体', 177800, True)
        self.document.tables[0].cell(2, 3).text = str(round(data[0]['sum'].sum(), 2))
        self.style_cell(self.document.tables[0].cell(2, 3), '宋体', 177800)
        self.document.tables[7].cell(len(data[0]) + 1, 2).text = str(round(data[0]['sum'].sum(), 2))
        self.style_cell(self.document.tables[7].cell(len(data[0]) + 1, 2), '宋体', 177800, True)
        self.document.tables[0].cell(2, 4).text = str(round(data[0]['max'].max(), 2))
        self.style_cell(self.document.tables[0].cell(2, 4), '宋体', 177800)
        self.document.tables[7].cell(len(data[0]) + 1, 3).text = str(round(data[0]['max'].max(), 2))
        self.style_cell(self.document.tables[7].cell(len(data[0]) + 1, 3), '宋体', 177800, True)
        self.document.tables[0].cell(2, 5).text = str(round(data[0]['min'].min(), 2))
        self.style_cell(self.document.tables[0].cell(2, 5), '宋体', 177800)
        self.document.tables[7].cell(len(data[0]) + 1, 3).text = str(round(data[0]['min'].min(), 2))
        self.style_cell(self.document.tables[7].cell(len(data[0]) + 1, 3), '宋体', 177800, True)
        for x in range(len(data[0])):
            self.document.tables[8].cell(1 + x, 0).text = data[0].index[x][0]
            self.style_cell(self.document.tables[8].cell(1 + x, 0), '宋体', 177800)
            self.document.tables[8].cell(1 + x, 1).text = data[0].index[x][1]
            self.style_cell(self.document.tables[8].cell(1 + x, 1), '宋体', 177800)
            self.document.tables[8].cell(1 + x, 2).text = data[0].loc[data[0].index[x], 'count']
            self.style_cell(self.document.tables[8].cell(1 + x, 2), '宋体', 177800)
            self.document.tables[8].cell(1 + x, 3).text = data[0].loc[data[0].index[x], 'sum']
            self.style_cell(self.document.tables[8].cell(1 + x, 3), '宋体', 177800)
            self.document.tables[8].cell(1 + x, 4).text = data[0].loc[data[0].index[x], 'max']
            self.style_cell(self.document.tables[8].cell(1 + x, 4), '宋体', 177800)
            self.document.tables[8].cell(1 + x, 5).text = data[0].loc[data[0].index[x], 'min']
            self.style_cell(self.document.tables[8].cell(1 + x, 5), '宋体', 177800)
        for x in range(len(data[1])):
            self.document.tables[9].cell(1 + x, 0).text = data[1].index[x][0]
            self.style_cell(self.document.tables[9].cell(1 + x, 0), '宋体', 177800)
            self.document.tables[9].cell(1 + x, 1).text = data[1].index[x][1]
            self.style_cell(self.document.tables[9].cell(1 + x, 1), '宋体', 177800)
            self.document.tables[9].cell(1 + x, 2).text = data[1].loc[data[1].index[x], 'count']
            self.style_cell(self.document.tables[9].cell(1 + x, 2), '宋体', 177800)
            self.document.tables[9].cell(1 + x, 3).text = data[1].loc[data[1].index[x], 'sum']
            self.style_cell(self.document.tables[9].cell(1 + x, 3), '宋体', 177800)
            self.document.tables[9].cell(1 + x, 4).text = data[1].loc[data[1].index[x], 'max']
            self.style_cell(self.document.tables[9].cell(1 + x, 4), '宋体', 177800)
            self.document.tables[9].cell(1 + x, 5).text = data[1].loc[data[1].index[x], 'min']
            self.style_cell(self.document.tables[9].cell(1 + x, 5), '宋体', 177800)

        data = Party(ty.stream.stream, lc.stream.stream, bank.date).get()
        self.sharp_table(self.document.tables[10], len(data) + 4)
        for x in range(len(data)):
            self.document.tables[10].cell(2 + x, 0).text = data.index[x][0]
            self.style_cell(self.document.tables[10].cell(2 + x, 0), '宋体', 177800)
            self.document.tables[10].cell(2 + x, 1).text = data.index[x][1]
            self.style_cell(self.document.tables[10].cell(2 + x, 1), '宋体', 177800)
            self.document.tables[10].cell(2 + x, 2).text = data.index[x][2]
            self.style_cell(self.document.tables[10].cell(2 + x, 2), '宋体', 177800)
            self.document.tables[10].cell(2 + x, 3).text = data.loc[data.index[x], 'sum']
            self.style_cell(self.document.tables[10].cell(2 + x, 3), '宋体', 177800)
            self.document.tables[10].cell(2 + x, 4).text = data.loc[data.index[x], 'count']
            self.style_cell(self.document.tables[10].cell(2 + x, 4), '宋体', 177800)
            self.document.tables[10].cell(2 + x, 5).text = data.loc[data.index[x], 'max']
            self.style_cell(self.document.tables[10].cell(2 + x, 5), '宋体', 177800)
            self.document.tables[10].cell(2 + x, 6).text = data.loc[data.index[x], 'min']
            self.style_cell(self.document.tables[10].cell(2 + x, 6), '宋体', 177800)

        if os.path.exists('风险管理部金融市场风险监测报告.docx'):
            os.remove('风险管理部金融市场风险监测报告.docx')
        self.document.save('风险管理部金融市场风险监测报告.docx')


if __name__ == '__main__':
    Word().go()
